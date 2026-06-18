# -*- coding: utf-8 -*-
"""
주식 단타 실전 교육 시스템 — PRD 자동 생성기
- 기획서 v1(주식단타교육_기획서_v1_20260527.docx) + 구현된 단타보조시스템 기반
- 포맷 규칙: CLAUDE_음성변환학습.md <word_format> 준수 (굴림 / 네이비·블루 / 명사형 종결 / 머리꼬리말 제거)
- make_plan.py 의 헬퍼와 동일 포맷 사용 (기획서와 시각적 일관성 유지)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "굴림"
COLOR_NAVY = RGBColor(0x00, 0x33, 0x66)
COLOR_BLUE = RGBColor(0x00, 0x70, 0xC0)
COLOR_TH = RGBColor(0x1F, 0x38, 0x64)
COLOR_TR_EVEN = RGBColor(0xD6, 0xE4, 0xF0)
COLOR_RED = RGBColor(0xFF, 0x00, 0x00)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, rgb):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_WHITE
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "003366")
        pPr.append(shd)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "200")
        pPr.append(ind)
    else:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLOR_BLUE
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    return p


def add_body(doc, text, bold=False, red=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(11)
    run.font.bold = bold
    if red:
        run.font.color.rgb = COLOR_RED
    p.paragraph_format.line_spacing = Pt(11 * 1.3)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(f"• {text}")
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(11)
    p.paragraph_format.line_spacing = Pt(11 * 1.3)
    p.paragraph_format.left_indent = Cm(0.5)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], COLOR_TH)
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.name = FONT
                run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = COLOR_WHITE
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            if ri % 2 == 1:
                set_cell_bg(cells[ci], COLOR_TR_EVEN)
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    run.font.name = FONT
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
                    run.font.size = Pt(10)
    return table


doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(0)
    section.footer_distance = Cm(0)
    for hp in section.header.paragraphs:
        hp.clear()
    for fp in section.footer.paragraphs:
        fp.clear()

# ── 표지 ──────────────────────────────────────
for _ in range(3):
    doc.add_paragraph()
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("주식 단타 실전 교육 시스템")
r.font.name = FONT; r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
r.font.size = Pt(24); r.font.bold = True; r.font.color.rgb = COLOR_NAVY
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sp.add_run("제품 요구사항 정의서 (PRD)")
r2.font.name = FONT; r2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
r2.font.size = Pt(14); r2.font.color.rgb = COLOR_BLUE
sp2 = doc.add_paragraph(); sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2b = sp2.add_run("테마주 단타 교육 커리큘럼 + 단타 보조 시스템 통합 명세")
r2b.font.name = FONT; r2b._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
r2b.font.size = Pt(11); r2b.font.color.rgb = COLOR_BLUE
for _ in range(2):
    doc.add_paragraph()
dp = doc.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = dp.add_run("문서 버전: v1   |   작성일: 2026년 06월 16일   |   기준 기획서: 주식단타교육_기획서_v1_20260527")
r3.font.name = FONT; r3._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
r3.font.size = Pt(11)
doc.add_page_break()

# ── 1. Product Overview ───────────────────────
add_heading(doc, "1. Product Overview")
add_heading(doc, "1.1 제품 정의", 2)
add_body(doc, "테마주 단타 매매를 실전 중심으로 학습하는 교육 커리큘럼과, 학습·실전을 동시에 지원하는 로컬 웹 기반 단타 보조 시스템으로 구성된 통합 제품.")
add_body(doc, "HTS(영웅문) 세팅·장 흐름 이해·테마주 대장주 전략·시간외 매매·리스크 관리·실전 일지의 6단계 교육과, KIS Open API 연동 실시간 급등 랭킹·상한가 체크리스트·강일지·관심종목 모니터링 도구가 하나의 학습 루프로 연결된 형태.")
add_heading(doc, "1.2 배경 및 문제 정의", 2)
add_bullet(doc, "시중 주식 교육의 차트 이론 편중 → 실전 단타 적용 난이도 높음")
add_bullet(doc, "HTS 실전 조작·호가창 독법·시간외 거래의 체계적 교육 부재")
add_bullet(doc, "실전 일지 기반 피드백 부재 → 매매 패턴 개선 속도 저하")
add_bullet(doc, "교육 내용을 즉시 적용·기록할 도구 부재 → 이론과 실전의 단절")
add_heading(doc, "1.3 핵심 가치 제안", 2)
add_bullet(doc, "라이브 실전 화면 공유 기반 100% 실전 중심 교육")
add_bullet(doc, "단계별 커리큘럼(세팅 → 흐름 → 전략 → 일지)과 보조 도구의 결합")
add_bullet(doc, "강일지 자동 기록·손익 추적으로 학습자 성과의 정량적 가시화")
add_bullet(doc, "비공개 소규모 집중 교육으로 질 높은 1:1 피드백 제공")
add_heading(doc, "1.4 제품 범위", 2)
add_table(doc, ["구분", "포함(In-Scope)", "제외(Out-of-Scope)"], [
    ["교육", "6단계 커리큘럼·실전 라이브·강일지 작성 지도", "투자 자문·종목 추천·수익 보장"],
    ["도구", "급등 랭킹·체크리스트·강일지·관심종목 모니터링", "자동 매매·주문 실행·자산 이체"],
    ["데이터", "KIS Open API 시세·체결·랭킹 조회", "외부 유료 데이터·뉴스 크롤링"],
])
add_body(doc, "※ 본 제품은 교육·기록 보조 목적이며, 매매 의사결정과 그 결과의 책임은 전적으로 사용자에게 귀속됨.", red=True)
doc.add_page_break()

# ── 2. 목표 사용자 (Persona) ───────────────────
add_heading(doc, "2. 목표 사용자 (Persona)")
add_heading(doc, "2.1 핵심 타겟", 2)
add_table(doc, ["구분", "특성"], [
    ["주요 타겟", "HTS 기본 조작은 가능하나 실전 단타 경험이 부족한 입문자~중급자"],
    ["연령대", "20~40대 직장인 또는 전업 투자 지망생"],
    ["자본 규모", "실전 매매 자금 500만~5,000만 원 보유 가능 계층"],
    ["학습 목표", "상한가 매매·테마주 단타를 통한 월 안정 수익 실현"],
])
add_heading(doc, "2.2 페르소나", 2)
add_body(doc, "페르소나 A — 직장인 입문자", bold=True)
add_bullet(doc, "상황: HTS 설치는 했으나 화면 구성·주문 속도 미숙, 뉴스 발생 후 대응 지연")
add_bullet(doc, "니즈: 영웅문 화면 세팅 표준안·아침 루틴·진입 타이밍 판단 기준")
add_bullet(doc, "기대: 보조 도구의 급등 랭킹·체크리스트로 진입 판단 시간 단축")
add_body(doc, "페르소나 B — 전업 지망 중급자", bold=True)
add_bullet(doc, "상황: 매매 경험은 있으나 손절 미준수·상한가 풀림 대응 부재로 손실 누적")
add_bullet(doc, "니즈: 강일지 기반 패턴 복기·금요일 리스크 관리·자금 관리 원칙")
add_bullet(doc, "기대: 강일지 손익 자동 집계로 자기 매매 패턴의 객관적 진단")
add_heading(doc, "2.3 니즈 및 페인포인트", 2)
add_bullet(doc, "HTS 화면 구성 미숙으로 시장 반응 속도 저하")
add_bullet(doc, "테마주 대장·부대장 구분 기준 불명확으로 진입 타이밍 실패 반복")
add_bullet(doc, "상한가 진입 후 풀림 시 대응 전략 부재로 손실 확대")
add_bullet(doc, "실전 일지 작성 습관 부재로 매매 패턴 개선 정체")
add_bullet(doc, "시간외 거래 전략(보통주·우선주)에 대한 체계적 이해 부족")
doc.add_page_break()

# ── 3. 핵심 기능 요구사항 ──────────────────────
add_heading(doc, "3. 핵심 기능 요구사항")
add_body(doc, "우선순위 — P0: 필수(MVP) / P1: 중요 / P2: 선택", bold=True)
add_table(doc, ["ID", "기능", "요구사항 요약", "우선순위"], [
    ["FR-1", "실시간 급등 랭킹", "코스피·코스닥 등락률 상위 30종목, 5초 주기 갱신, 20%+ 강조", "P0"],
    ["FR-2", "상한가 체크리스트", "6항목 점수화(0~100)로 진입검토/관망/제외 자동 판정", "P0"],
    ["FR-3", "강일지(매매 일지)", "매수·매도 기록, 손익 자동 계산, 일자별 영구 저장", "P0"],
    ["FR-4", "관심종목 실시간", "최대 20종목 실시간 체결가, 목표가·손절가 도달 알림", "P0"],
    ["FR-5", "장 흐름 자동 구분", "장중(09:00~15:30)에만 폴링, 장마감 시 자동 중단", "P1"],
    ["FR-6", "최근 7일 손익 추이", "강일지 일자별 손익 미니 차트 시각화", "P1"],
    ["FR-7", "교육 커리큘럼 연동", "강의 단계별 산출물(세팅·모의/실전 강일지) 점검", "P2"],
])
add_heading(doc, "3.1 상한가 체크리스트 점수 규칙", 2)
add_body(doc, "6개 항목 가중 합산(만점 100점), 70점 이상 진입검토 · 50~69점 관망 · 50점 미만 제외.")
add_table(doc, ["항목", "판정 기준", "배점"], [
    ["당일 호재 뉴스", "당일 강력 호재(자율주행·전기차·정치 테마 등) 존재", "30"],
    ["거래량 3배 이상", "전일 대비 거래량 비율 ≥ 3.0", "20"],
    ["등락률 20% 이상", "당일 등락률 ≥ 20%", "20"],
    ["네이버 실검 등재", "네이버 실시간 검색 등재 여부", "15"],
    ["시총 1,000억 이하", "시가총액 소형주 조건 충족", "10"],
    ["연속 상한가 아님", "당일 등락률 < 29.5%(2연상 이력 회피)", "5"],
])
doc.add_page_break()

# ── 4. 비기능 요구사항 ─────────────────────────
add_heading(doc, "4. 비기능 요구사항")
add_table(doc, ["분류", "요구사항"], [
    ["성능", "랭킹 폴링 5초 주기, 관심종목 체결가는 WebSocket 실시간 수신"],
    ["가용성", "WebSocket 단절 시 5초 후 자동 재연결 및 구독 복구"],
    ["안정성", "KIS 키 미설정·API 실패 시에도 서버 비중단(폴링만 비활성)"],
    ["보안", "API 키는 .env 로컬 보관, 저장소(git) 추적 제외, 토큰 24시간 캐시"],
    ["이식성", "프론트엔드는 외부 CDN 없는 단일 HTML, 오프라인 UI 로드 가능"],
    ["영속성", "강일지·관심종목은 SQLite(trades.db)에 영구 저장"],
    ["운영", "로컬 단일 명령(python main.py) 기동, 브라우저 localhost 접속"],
])
add_body(doc, "※ 실거래 주문·자금 이체 기능 미포함 — 시세 조회·기록 전용으로 한정해 오작동 위험 최소화.")
doc.add_page_break()

# ── 5. 사용자 스토리 ───────────────────────────
add_heading(doc, "5. 사용자 스토리")
add_table(doc, ["ID", "사용자 스토리", "수용 기준"], [
    ["US-1", "입문자로서 급등 종목을 한눈에 확인 원함", "랭킹 패널에서 등락률 상위 종목·거래량비 표시, 20%+ 강조"],
    ["US-2", "입문자로서 진입 가치를 빠르게 판단 원함", "종목 클릭 시 체크리스트 자동 입력·점수·판정 즉시 산출"],
    ["US-3", "중급자로서 매매를 기록·복기 원함", "매수·매도 입력 시 손익 자동 계산·당일 합계·7일 추이 표시"],
    ["US-4", "중급자로서 목표·손절 도달을 즉시 인지 원함", "관심종목 실시간 가격이 목표가/손절가 도달 시 알림 발생"],
    ["US-5", "학습자로서 장 시간에만 데이터 집중 원함", "09:00~15:30 외에는 폴링 중단으로 불필요한 호출 차단"],
])
doc.add_page_break()

# ── 6. 기능 명세 (Feature Spec) ────────────────
add_heading(doc, "6. 기능 명세 (Feature Spec)")
add_heading(doc, "6.1 시스템 구성", 2)
add_body(doc, "FastAPI 백엔드가 KIS REST를 5초 주기로 폴링해 랭킹을 갱신하고, 관심종목은 KIS WebSocket으로 실시간 체결가 수신. 브라우저는 SSE로 갱신 데이터를 받으며, 강일지는 SQLite에 영구 저장. 프론트엔드는 외부 CDN 없는 단일 HTML 4패널 대시보드.")
add_table(doc, ["모듈", "역할"], [
    ["main.py", "FastAPI 앱 진입점·라우터·lifespan(스케줄러/WS 기동)"],
    ["kis_client.py", "KIS REST 토큰 캐시·랭킹·현재가 조회"],
    ["ws_manager.py", "KIS WebSocket 구독·자동 재연결"],
    ["sse_broker.py", "브라우저 실시간 푸시(SSE) 브로드캐스트"],
    ["scheduler.py", "APScheduler 5초 랭킹 폴링(장중 한정)"],
    ["database.py", "SQLite 모델·CRUD(trades/watchlist/checklist)"],
    ["static/index.html", "4패널 단일 페이지(랭킹/체크리스트/강일지/관심종목)"],
])
add_heading(doc, "6.2 API 엔드포인트", 2)
add_table(doc, ["메서드·경로", "기능"], [
    ["GET /", "대시보드 페이지 반환"],
    ["GET /events", "SSE 실시간 이벤트 스트림(ranking/price_update/ws_status)"],
    ["GET /api/ranking?market=", "캐시된 급등 랭킹 반환(J/Q/ALL)"],
    ["POST /api/trades", "강일지 매매 기록 저장"],
    ["GET /api/trades?date=", "일자별 매매 기록 조회"],
    ["DELETE /api/trades/{id}", "매매 기록 삭제"],
    ["POST·GET /api/watchlist", "관심종목 등록·목록 조회(최대 20종목)"],
    ["DELETE /api/watchlist/{ticker}", "관심종목 해제 및 WS 구독 해지"],
    ["POST /api/checklist", "체크리스트 점수 산출·이력 저장"],
])
add_heading(doc, "6.3 데이터 모델 (SQLite)", 2)
add_table(doc, ["테이블", "주요 컬럼"], [
    ["trades", "date·ticker·name·side(BUY/SELL)·price·amount·reason·pnl"],
    ["watchlist", "ticker(UNIQUE)·name·target_price·stop_loss·memo"],
    ["checklist_history", "ticker·date·score·news/volume/surge/naver/cap/no_chain·result"],
])
add_heading(doc, "6.4 화면 구성 (4패널)", 2)
add_table(doc, ["패널", "내용"], [
    ["급등 랭킹", "코스피/코스닥/전체 탭, 등락률·거래량비, 20%+ 강조, 클릭 시 체크리스트 연동"],
    ["상한가 체크리스트", "종목코드 입력·자동입력, 6항목 점수·판정 표시"],
    ["강일지", "매수/매도 모달, 손익 자동 계산, 당일 합계, 7일 손익 미니 차트"],
    ["관심종목", "실시간 체결가, 목표가/손절가 도달 시 토스트·브라우저 알림"],
])
doc.add_page_break()

# ── 7. 제약 조건 및 가정 ───────────────────────
add_heading(doc, "7. 제약 조건 및 가정")
add_heading(doc, "7.1 제약 조건", 2)
add_bullet(doc, "KIS Open API 키(앱키·시크릿·계좌번호) 발급 및 .env 설정 필요")
add_bullet(doc, "KIS API 호출 한도·장 운영 시간(09:00~15:30) 의존")
add_bullet(doc, "로컬 실행 환경(Python 3.11+ 권장) 및 브라우저 필요")
add_bullet(doc, "실거래 주문·자금 이동 기능 미포함(시세 조회·기록 전용)")
add_heading(doc, "7.2 가정", 2)
add_bullet(doc, "사용자는 영웅문 등 실거래 HTS를 별도 보유·운용")
add_bullet(doc, "테마주(정치·산업 이슈)의 단기 급등락 기회 상시 존재")
add_bullet(doc, "학습자는 소액 시뮬레이션 후 비중을 단계적으로 확대")
add_heading(doc, "7.3 리스크 및 대응", 2)
add_table(doc, ["리스크", "대응 방안"], [
    ["학습자 손실 시 심리적 이탈", "초기 소액(50만~100만 원) 시뮬레이션 후 비중 단계 확대"],
    ["상한가 풀림 손실", "금요일 포지션 경량화·풀림 즉시 손절 원칙 반복 훈련"],
    ["미확인 뉴스 오매매", "뉴스 끝까지 확인(기업 측 답변 포함) 습관화"],
    ["API 장애·키 미설정", "폴링 자동 비활성·서버 비중단·재연결로 가용성 확보"],
])
doc.add_page_break()

# ── 8. 성공 지표 ───────────────────────────────
add_heading(doc, "8. 성공 지표 (KPI)")
add_heading(doc, "8.1 학습 성과 지표", 2)
add_table(doc, ["지표", "목표치", "측정 방법"], [
    ["상한가 포착률", "5종목 중 3종목 이상(60%)", "강일지 성공/실패 건수 집계"],
    ["월 수익률", "전체 자금 대비 +5~10%", "강일지 손익 합계"],
    ["손절 준수율", "풀림 즉시 손절 90% 이상", "손절 일지 기록 여부"],
    ["HTS 대응 속도", "뉴스 발생 후 30초 내 진입", "타임스탬프 기록"],
])
add_heading(doc, "8.2 제품 사용 지표", 2)
add_table(doc, ["지표", "목표치"], [
    ["강일지 작성 지속률", "수강 기간 중 거래일 80% 이상 기록"],
    ["체크리스트 활용률", "진입 전 체크리스트 실행 비율 70% 이상"],
    ["보조 도구 가동률", "장중 대시보드 상시 가동(연결 유지율 95% 이상)"],
])
doc.add_page_break()

# ── 9. 미결 이슈 ───────────────────────────────
add_heading(doc, "9. 미결 이슈 (Open Issues)")
add_table(doc, ["ID", "이슈", "현재 상태 / 검토 방향"], [
    ["OI-1", "뉴스 호재 강도 자동 판정", "현재 수동 입력(confirm) → 뉴스 API 연동 자동화 검토(P2)"],
    ["OI-2", "네이버 실검 등재 자동 확인", "현재 수동 입력 → 외부 데이터 연동 시 자동화 가능성 검토"],
    ["OI-3", "손익 계산 정합성", "당일 LIFO 매칭 기준 → 분할 매수·익절 케이스 정밀화 필요"],
    ["OI-4", "체크리스트 가중치 튜닝", "현 배점(30/20/20/15/10/5) → 실전 데이터 누적 후 보정 검토"],
    ["OI-5", "교육 LMS 연동", "강의 일정·산출물 관리의 별도 LMS 통합 여부 미정"],
])
add_body(doc, "")
add_body(doc, "— 본 PRD는 기획서 v1과 구현된 단타 보조 시스템(테스트 9/9 통과·E2E 동작 확인)을 기준으로 작성된 v1 문서이며, 실전 데이터 누적과 교육 운영 피드백에 따라 개정 예정.")

import os
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "주식단타교육_PRD_v1_20260616.docx")
doc.save(out)
print("저장 완료:", out)
