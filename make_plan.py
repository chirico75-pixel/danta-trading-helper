from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

FONT = "굴림"
COLOR_NAVY = RGBColor(0x00, 0x33, 0x66)
COLOR_BLUE = RGBColor(0x00, 0x70, 0xC0)
COLOR_TH = RGBColor(0x1F, 0x38, 0x64)
COLOR_TR_EVEN = RGBColor(0xD6, 0xE4, 0xF0)
COLOR_RED = RGBColor(0xFF, 0x00, 0x00)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)

def p_font(paragraph, size, bold=False, color=None):
    for run in paragraph.runs:
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_WHITE
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "003366")
        pPr.append(shd)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        pPr2 = p._p.pPr
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "200")
        pPr2.append(ind)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLOR_BLUE
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    return p

def add_body(doc, text, bold=False, red=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(11)
    run.font.bold = bold
    if red:
        run.font.color.rgb = COLOR_RED
    p.paragraph_format.line_spacing = Pt(11 * 1.3)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(f"• {text}")
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(11)
    p.paragraph_format.line_spacing = Pt(11 * 1.3)
    p.paragraph_format.left_indent = Cm(0.5)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], COLOR_TH)
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.name = FONT
                run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = COLOR_WHITE
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            if ri % 2 == 1:
                set_cell_bg(cells[ci], COLOR_TR_EVEN)
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    run.font.name = FONT
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
                    run.font.size = Pt(10)
    return table

doc = Document()

# 여백 설정
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(0)
    section.footer_distance = Cm(0)
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True
    for hdr_para in section.header.paragraphs:
        hdr_para.clear()
    for ftr_para in section.footer.paragraphs:
        ftr_para.clear()

# ─────────────────────────────────────────────
# 표지
# ─────────────────────────────────────────────
for _ in range(3):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("주식 단타 실전 교육 강의")
r.font.name = FONT
r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
r.font.size = Pt(24)
r.font.bold = True
r.font.color.rgb = COLOR_NAVY

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run("테마주 단타 매매 & HTS 실전 운용 시스템 기획서")
r2.font.name = FONT
r2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
r2.font.size = Pt(14)
r2.font.color.rgb = COLOR_BLUE

for _ in range(2):
    doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = date_p.add_run("작성일: 2026년 05월 27일")
r3.font.name = FONT
r3._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
r3.font.size = Pt(11)

doc.add_page_break()

# ─────────────────────────────────────────────
# 1. 기획 개요
# ─────────────────────────────────────────────
add_heading(doc, "1. 기획 개요")

add_heading(doc, "1.1 프로젝트명 및 목적", level=2)
add_body(doc, "프로젝트명: 주식 단타 실전 교육 강의 시스템")
add_body(doc, "목적: 영웅문 HTS 실전 세팅부터 테마주 상한가 매매까지, 실전 단타 매매의 전 과정을 라이브 강의 및 실전 일지 기반으로 체계적으로 전달하는 비공개 1:1 실전 교육 프로그램 구성 및 운영")

add_heading(doc, "1.2 기획 배경 및 문제 정의", level=2)
add_bullet(doc, "시중에 유통되는 대부분의 주식 교육은 차트 이론 위주로 실전 단타 매매에 직접 적용하기 어려움")
add_bullet(doc, "HTS(영웅문) 세팅, 호가창 독법, 시간외 거래 등 실전 조작 방법을 체계적으로 다루는 교육 부재")
add_bullet(doc, "상한가 매매, 테마주 순환 매매 등 고수익 전략은 실전 경험 없이는 습득 불가")
add_bullet(doc, "실전 일지 기반의 피드백 시스템 부재로 학습자 성과 개선 속도 저하")

add_heading(doc, "1.3 핵심 가치 제안", level=2)
add_bullet(doc, "라이브 실전 화면 공유를 통한 100% 실전 중심 교육")
add_bullet(doc, "단계별 커리큘럼(HTS 세팅 → 장 흐름 이해 → 테마주 전략 → 실전 일지)")
add_bullet(doc, "실시간 매매 일지(강일지) 작성 시스템으로 학습자 성과 추적")
add_bullet(doc, "비공개 소규모 집중 교육으로 질 높은 1:1 피드백 제공")

doc.add_paragraph()

# ─────────────────────────────────────────────
# 2. 현황 분석
# ─────────────────────────────────────────────
add_heading(doc, "2. 현황 분석")

add_heading(doc, "2.1 시장·환경 현황", level=2)
add_bullet(doc, "국내 주식 시장 일평균 거래대금 약 10~20조 원 수준 유지, 개인 투자자 비중 지속 증가")
add_bullet(doc, "2020~2021년 동학개미운동 이후 개인 투자자의 단기매매 수요 급증")
add_bullet(doc, "테마주(정치·이슈 테마)의 단기 급등락 반복으로 단타 매매 기회 상시 존재")
add_bullet(doc, "코스피·코스닥 시장 내 상한가(+30%) 종목 일평균 2~10개 발생")

add_heading(doc, "2.2 기회 및 위협 요인", level=2)
add_table(doc,
    ["구분", "내용"],
    [
        ["기회", "정치 테마주(대선·총선), 산업 이슈 테마(요소수·자율주행·메타버스 등) 지속 발생"],
        ["기회", "시간외 거래를 통한 다음날 시초가 예측 전략 유효성 높음"],
        ["기회", "영웅문 HTS 커스터마이징으로 실시간 대응 속도 최적화 가능"],
        ["위협", "상한가 풀림(3시 20분 이후 집중) 시 예상치 못한 손실 발생 가능"],
        ["위협", "금요일장 특성상 상한가 풀림 빈도 증가로 리스크 집중"],
        ["위협", "미확인·과장 뉴스로 인한 오매매 위험 상존"],
    ]
)

doc.add_paragraph()

# ─────────────────────────────────────────────
# 3. 타겟 및 수요 분석
# ─────────────────────────────────────────────
add_heading(doc, "3. 타겟 및 수요 분석")

add_heading(doc, "3.1 핵심 타겟 정의", level=2)
add_table(doc,
    ["구분", "특성"],
    [
        ["주요 타겟", "HTS 기본 조작은 가능하나 실전 단타 매매 경험이 부족한 주식 입문자~중급자"],
        ["연령대", "20~40대 직장인 또는 전업 투자 지망생"],
        ["자본 규모", "실전 매매 자금 500만~5,000만 원 보유 가능 계층"],
        ["학습 목표", "상한가 매매·테마주 단타를 통한 월 안정 수익 실현"],
    ]
)

add_heading(doc, "3.2 학습자 니즈 및 페인포인트", level=2)
add_bullet(doc, "HTS 화면 구성 방법을 몰라 시장 반응 속도가 느림")
add_bullet(doc, "테마주 대장주·부대장주 구분 기준 불명확으로 진입 타이밍 실패 반복")
add_bullet(doc, "상한가 진입 후 풀림 시 대응 전략 부재로 손실 확대")
add_bullet(doc, "실전 일지 작성 습관 없어 매매 패턴 개선이 이루어지지 않음")
add_bullet(doc, "시간외 거래 전략(보통주·우선주 활용)에 대한 체계적 이해 부족")

doc.add_paragraph()

# ─────────────────────────────────────────────
# 4. 프로그램 구성
# ─────────────────────────────────────────────
add_heading(doc, "4. 프로그램·서비스 구성")

add_heading(doc, "4.1 강의 커리큘럼 구성", level=2)
add_table(doc,
    ["단계", "강의 번호", "주제", "핵심 내용"],
    [
        ["1단계\n기초 세팅", "1-1, 1-2, 1-3", "영웅문 HTS 세팅", "화면 구성·관심종목·주문창·미수·신용 설정법"],
        ["2단계\n장 흐름", "1-5, 1-6, 1-7, 1-8", "하루 장 흐름 완전 이해", "장전 동시호가·장중 흐름·시간외 매매·장마감 처리"],
        ["3단계\n테마주 전략", "2, 3, 5", "테마주 대장·부대장 전략", "이슈 테마 발굴·대장주 진입·순환매 타이밍·상한가 포착"],
        ["4단계\n특수 매매", "4-1, 4-2, 4-3", "시간외·미수거래", "보통주·우선주 시간외 전략·미수거래 리스크 관리"],
        ["5단계\n리스크 관리", "6", "금요일장 위험·상한가 풀림", "장 마지막 20분 풀림 대응·재진입 판단 기준"],
        ["6단계\n실전 일지", "실전 라이브", "강일지·실전 일지 작성", "3천만 원 이상 매매 기록·시초가 대응·성과 추적"],
    ]
)

add_heading(doc, "4.2 핵심 강의 내용 요약", level=2)

add_body(doc, "▶ HTS 세팅 핵심", bold=True)
add_bullet(doc, "화면 2분할: 코스피·코스닥 전일대비 상승률 관심종목(4번)·코스피 5번 관찰")
add_bullet(doc, "퀵 주문창 연동 설정 및 시장가/지정가 전환 방법")
add_bullet(doc, "예상 체결 등락 메뉴 등록 및 아침 20분 예상 등락 분석 루틴")
add_bullet(doc, "미국 지수 선물(나스닥·S&P500) 실시간 확인 화면 구성")

add_body(doc, "▶ 테마주 매매 핵심 원칙", bold=True)
add_bullet(doc, "대장주만 매매: 부대장·군소 테마주 진입 금지 → 상승폭 및 유동성 차이")
add_bullet(doc, "호재 뉴스 강도 판단: 절대적 강도(뉴스 내용) × 상대적 강도(전날·당일 거래량 비교)")
add_bullet(doc, "순환매 패턴 활용: 한 테마가 끝나면 연관 테마로 이동하는 흐름 추종")
add_bullet(doc, "상한가 풀림 대응: 풀림 후 재진입 시 손실 누적 패턴 회피 → 즉시 손절 원칙")

add_body(doc, "▶ 시간외 거래 전략", bold=True)
add_bullet(doc, "시간외 매도 호가(3:40~4:00): 상한가 물량 배정 순서 최우선 확보")
add_bullet(doc, "우선주 활용: 보통주 대비 시가총액 소규모 → 변동폭 극대화 가능")
add_bullet(doc, "시간외 매수 후 다음날 시초가 플러스 시작 시 즉시 매도 전략")

add_body(doc, "▶ 금요일장 특별 주의사항", bold=True, red=True)
add_bullet(doc, "금요일 3시 이후 상한가 풀림 빈도 평일 대비 높음 → 보유 포지션 경량화")
add_bullet(doc, "3일 이상 된 뉴스 재부각 테마는 당일 호재 없는 경우 풀림 위험 증가")
add_bullet(doc, "미확인 뉴스(기업 측 '확인 불가' 답변) 기반 상한가는 당일 풀림 확률 높음")

add_heading(doc, "4.3 실전 일지 시스템 (강일지)", level=2)
add_bullet(doc, "3,000만 원 이상 매매 건에 대해 실시간 일지 기록")
add_bullet(doc, "기록 항목: 진입 이유(호재 뉴스) / 매수 단가 및 금액(1차~6차) / 매도 타이밍 / 손익")
add_bullet(doc, "당일 손실 매도 사례도 별도 기록하여 패턴 개선에 활용")
add_bullet(doc, "시초가 방향성 확인 매도 여부 기록 → 다음날 전략 수립에 연결")

doc.add_paragraph()

# ─────────────────────────────────────────────
# 5. 전략 및 방법론
# ─────────────────────────────────────────────
add_heading(doc, "5. 전략 및 방법론")

add_heading(doc, "5.1 단타 매매 핵심 프레임워크", level=2)
add_table(doc,
    ["단계", "행동", "판단 기준"],
    [
        ["장 전(8:00~9:00)", "예상 등락 분석·시간외 잔고 점검", "전날 상한가 종목 시간외 체결량·예상 등락 상위 확인"],
        ["장 시작(9:00~)", "동시호가 진입 여부 결정", "뉴스 강도·예상 등락 위치·코스피 선물 방향"],
        ["장중(9:00~15:20)", "테마주 모니터링 및 진입·청산", "20% 이상 급등 종목 뉴스 확인·상대적 강도 비교"],
        ["장마감(15:20~15:30)", "상한가 물량 확인", "3:20 이후 풀림 여부·시간외 매도 호가 배치"],
        ["시간외(15:30~16:00)", "다음날 전략 종목 매수", "시간외 체결량·뉴스 지속성 판단"],
        ["장 후(16:00~)", "일지 작성 및 복기", "강일지 기록·다음날 주목 종목 관심목록 업데이트"],
    ]
)

add_heading(doc, "5.2 상한가 종목 선별 기준", level=2)
add_bullet(doc, "당일 강력 호재 뉴스 최우선: 자율주행·전기차·정치 테마 등 명확한 이슈")
add_bullet(doc, "오전 20% 이상 급등 후 네이버 실시간 인기 검색어 등재 여부 확인")
add_bullet(doc, "상대적 세기 비교: 동일 테마 내 다수 종목 중 거래량·뉴스 강도 1위 종목만 선택")
add_bullet(doc, "시가총액 조건: 소형주(500억 이하) 우선주는 변동성 극대화 가능하나 풀림 위험도 증가")
add_bullet(doc, "연속 상한가 이력: 2연상 이상 종목은 매도 압력 증가로 진입 자제")

add_heading(doc, "5.3 자금 관리 원칙", level=2)
add_bullet(doc, "1회 진입 금액: 500만 원 단위 시장가 매수 원칙")
add_bullet(doc, "최대 비중: 1개 종목 3,000만 원 이하 집중")
add_bullet(doc, "미수 거래: 3일 내 청산 조건으로 활용, 하락 시 즉시 손절")
add_bullet(doc, "신용 거래: 2~3개월 여유 가능하나 일반 단타 플레이에는 미사용")

doc.add_paragraph()

# ─────────────────────────────────────────────
# 6. 실행 계획
# ─────────────────────────────────────────────
add_heading(doc, "6. 실행 계획")

add_heading(doc, "6.1 강의 일정 및 마일스톤", level=2)
add_table(doc,
    ["단계", "기간", "주요 내용", "산출물"],
    [
        ["Phase 1", "1~2주차", "HTS 세팅·기본 주문 실습", "영웅문 화면 세팅 완료"],
        ["Phase 2", "3~4주차", "장 흐름 이해·예상 등락 분석 훈련", "아침 루틴 확립"],
        ["Phase 3", "5~7주차", "테마주 모의 매매·뉴스 강도 판단 훈련", "모의 강일지 10건"],
        ["Phase 4", "8~10주차", "실전 소규모 매매(100만~500만 원)", "실전 강일지 10건"],
        ["Phase 5", "11~12주차", "시간외·미수·상한가 전략 통합 실습", "종합 성과 리포트"],
    ]
)

add_heading(doc, "6.2 필요 자원", level=2)
add_bullet(doc, "HTS: 영웅문 S(키움증권) 설치 및 계좌 개설")
add_bullet(doc, "실전 자금: 최소 500만 원 이상 (Phase 4 진입 조건)")
add_bullet(doc, "강의 환경: 화면 공유 기반 라이브 강의 (Zoom/OBS)")
add_bullet(doc, "실전 일지 양식: 강일지 엑셀 템플릿 제공")

add_heading(doc, "6.3 리스크 및 대응", level=2)
add_table(doc,
    ["리스크", "대응 방안"],
    [
        ["학습자 손실 발생 시 심리적 이탈", "초기 소액(50만~100만 원) 시뮬레이션 후 비중 단계적 확대"],
        ["테마주 소멸 시 학습 기회 감소", "정치 이슈·산업 테마 항상 존재 → 이슈 발굴 훈련 병행"],
        ["상한가 풀림으로 인한 손실", "금요일 포지션 경량화·풀림 즉시 손절 원칙 반복 훈련"],
        ["미확인 뉴스 오매매", "뉴스 끝까지 확인(기업 측 답변 포함) 습관화"],
    ]
)

doc.add_paragraph()

# ─────────────────────────────────────────────
# 7. 기대 성과 및 KPI
# ─────────────────────────────────────────────
add_heading(doc, "7. 기대 성과 및 KPI")

add_heading(doc, "7.1 정량적 목표", level=2)
add_table(doc,
    ["지표", "목표치", "측정 방법"],
    [
        ["상한가 포착률", "5종목 중 3종목 이상 (60%)", "강일지 성공/실패 건수 집계"],
        ["월 수익률", "전체 자금 대비 +5~10%", "강일지 손익 합계"],
        ["손절 준수율", "풀림 즉시 손절 90% 이상", "손절 일지 기록 여부"],
        ["HTS 대응 속도", "뉴스 발생 후 30초 내 진입", "타임스탬프 기록"],
    ]
)

add_heading(doc, "7.2 정성적 목표", level=2)
add_bullet(doc, "테마주 흐름 직관적 판단 능력 체득 (6개월 반복 훈련 기준)")
add_bullet(doc, "강일지 자발적 작성 습관 확립")
add_bullet(doc, "상한가 풀림 및 금요일 리스크에 대한 반사적 대응 능력")
add_bullet(doc, "대장주·부대장주 즉각적 구분 능력")

doc.add_paragraph()

# ─────────────────────────────────────────────
# 부록: 주요 강의 내용 요약표
# ─────────────────────────────────────────────
add_heading(doc, "부록: 강의 파일 목록 및 내용 요약")

add_table(doc,
    ["파일명", "주제", "핵심 키워드"],
    [
        ["1-1(영웅문 셋팅1)", "HTS 화면 구성", "2화면 분할, 관심종목, 뉴스창, 예상 등락"],
        ["1-2(영웅문 셋팅2)", "주문창 세팅", "퀵주문, 시장가, 미수/신용, 금액 설정"],
        ["1-3(영웅문 셋팅3)", "HTS 세팅 완성", "단축키, 툴바 설정, 차트 분봉"],
        ["1-5(하루흐름2)", "장전 동시호가", "예상 체결가, 동시호가 진입 전략"],
        ["1-6(하루흐름3)", "장전 동시호가2", "아침 루틴, 예상 등락 분석"],
        ["1-7(하루흐름4)", "시간외 매매", "3:40~4:00 시간외 매도 물량 배정"],
        ["1-8(하루흐름5)", "장마감 처리", "상한가 잔고 처리, 다음날 전략 수립"],
        ["2(장마감,테마주 대장 부대장)", "테마주 대장 전략", "요소수 테마, 대장 교체 타이밍"],
        ["3(장마감,윤석열 테마)", "정치 테마 전략", "윤석열 테마주, 시간외 진입 사례"],
        ["4-1(시간외거래,보통주 우선주)", "보통주·우선주 시간외", "우선주 시가총액 활용, 뉴스 끝까지 확인"],
        ["4-2(시간외거래,보통주 우선주)", "시간외 실전 사례2", "연속 상한가, 시간외 물량 전략"],
        ["4-3(미수거래)", "미수거래 실전", "3일 내 청산, 미수 리스크 관리"],
        ["5(테마주 순환매)", "순환매 전략", "스남매(순환매) 패턴, 흐름 추종"],
        ["6(금요일장 위험성)", "금요일 리스크", "상한가 풀림 대응, 틀렸을 때 즉시 손절"],
        ["20211103~20211117", "실전 매매 일지", "효성오앤비 실전, 강일지 작성 시연"],
    ]
)

OUTPUT = r"G:\내 드라이브\Antigravity_Projects\기술적분석 강의\2301_진욱형님_비공개\주식단타교육_기획서_v1_20260527.docx"
doc.save(OUTPUT)
print(f"완료\n생성 파일: {OUTPUT}")
